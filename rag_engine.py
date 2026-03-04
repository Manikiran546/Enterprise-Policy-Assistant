"""
RAG Engine Module
Handles document processing, embedding, vector storage, and retrieval-augmented generation.
Uses API-based services only (no local ML models) for low-memory deployment.
"""

import os
import warnings
from collections import defaultdict

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEndpointEmbeddings
from langchain_groq import ChatGroq
from langchain_core.documents import Document

warnings.filterwarnings("ignore", category=FutureWarning)


class RAGEngine:
    """Per-session RAG pipeline: parse → chunk → embed → store → retrieve → generate."""

    # Shared across all sessions to save memory
    _embeddings = None
    _model = None

    @classmethod
    def get_embeddings(cls):
        if cls._embeddings is None:
            cls._embeddings = HuggingFaceEndpointEmbeddings(
                model="sentence-transformers/all-MiniLM-L6-v2",
                huggingfacehub_api_token=os.getenv("HUGGINGFACEHUB_API_TOKEN"),
            )
        return cls._embeddings

    @classmethod
    def get_model(cls):
        if cls._model is None:
            cls._model = ChatGroq(
                model="llama-3.3-70b-versatile",
                temperature=0.0,
                max_tokens=512,
            )
        return cls._model

    def __init__(self):
        self.vectorstore = None
        self.documents_loaded = []

    # ------------------------------------------------------------------ #
    #  Document Processing
    # ------------------------------------------------------------------ #

    def _extract_text_from_pdf(self, file_path: str) -> list[Document]:
        loader = PyPDFLoader(file_path)
        return loader.load()

    def _extract_text_from_docx(self, file_path: str, filename: str) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        docs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": filename, "page": i + 1},
                    )
                )
        return docs

    def _extract_text_from_txt(self, file_path: str, filename: str) -> list[Document]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return [
            Document(page_content=content, metadata={"source": filename, "page": 1})
        ]

    def process_documents(self, files: list[tuple[str, str, str]]) -> dict:
        all_docs = []

        for file_path, filename, ext in files:
            try:
                if ext == ".pdf":
                    docs = self._extract_text_from_pdf(file_path)
                    for doc in docs:
                        doc.metadata["source"] = filename
                elif ext == ".docx":
                    docs = self._extract_text_from_docx(file_path, filename)
                elif ext == ".txt":
                    docs = self._extract_text_from_txt(file_path, filename)
                else:
                    continue

                all_docs.extend(docs)
                self.documents_loaded.append(filename)
            except Exception as e:
                print(f"Error processing {filename}: {e}")
                continue

        if not all_docs:
            return {"status": "error", "message": "No documents could be processed."}

        # Chunk the documents
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=150,
        )
        chunks = splitter.split_documents(all_docs)

        for i, doc in enumerate(chunks):
            doc.metadata["id"] = i

        # Build or extend vector store
        if self.vectorstore is None:
            self.vectorstore = FAISS.from_documents(chunks, self.get_embeddings())
        else:
            self.vectorstore.add_documents(chunks)

        return {
            "status": "success",
            "chunks_created": len(chunks),
            "documents_processed": len(files),
            "filenames": [f[1] for f in files],
        }

    # ------------------------------------------------------------------ #
    #  RAG Query Pipeline
    # ------------------------------------------------------------------ #

    def query(self, user_query: str) -> dict:
        if self.vectorstore is None:
            return {
                "answer": "Please upload documents first before asking questions.",
                "sources": [],
            }

        retriever = self.vectorstore.as_retriever(
            search_type="mmr",
            search_kwargs={"k": 20, "lambda_mult": 0.5},
        )

        # -------- Query Rewrite --------
        rewrite_prompt = f"Rewrite this query for document retrieval:\n{user_query}"
        try:
            rewritten_query = self.get_model().invoke(rewrite_prompt).content.strip()
        except Exception:
            rewritten_query = user_query

        # -------- Multi Query (RAG Fusion) --------
        fusion_prompt = (
            f"Generate 3 search queries for:\n{rewritten_query}\n"
            "Only return queries newline separated."
        )
        try:
            response = self.get_model().invoke(fusion_prompt)
            queries = [q.strip() for q in response.content.strip().split("\n") if q.strip()]
            queries = queries[:3]
        except Exception:
            queries = [rewritten_query]

        # -------- Retrieval --------
        all_docs = []
        for q in queries:
            try:
                docs_q = retriever.invoke(q)
                all_docs.append(docs_q)
            except Exception:
                continue

        if not all_docs:
            return {
                "answer": "I encountered an error while searching the documents. Please try again.",
                "sources": [],
            }

        # -------- RAG Fusion Scoring --------
        fusion_scores = defaultdict(float)
        for docs_list in all_docs:
            for rank, doc in enumerate(docs_list):
                fusion_scores[doc.page_content] += 1 / (rank + 60)

        sorted_docs = sorted(
            fusion_scores.items(), key=lambda x: x[1], reverse=True
        )

        # Deduplicate and get top 5
        docs = []
        seen = set()
        for text, score in sorted_docs[:5]:
            if text in seen:
                continue
            seen.add(text)
            for dlist in all_docs:
                for d in dlist:
                    if d.page_content == text:
                        docs.append(d)
                        break
                else:
                    continue
                break

        if not docs:
            return {
                "answer": "The uploaded documents do not contain this information.",
                "sources": [],
            }

        # -------- Build Context --------
        context = "\n\n".join([d.page_content for d in docs])

        # -------- Generate Answer --------
        final_prompt = f"""You are a company policy assistant using RAG.

Answer the question using ONLY the provided context below.
If the context contains relevant information, provide a clear and helpful answer.
If the information is truly not present in the context at all, respond with:
"The uploaded documents do not contain this information."

Do not make up information not found in the context.
Do not use prior knowledge.

Context:
{context}

Question:
{user_query}
"""
        try:
            answer = self.get_model().invoke(final_prompt).content.strip()
        except Exception as e:
            answer = f"Error generating response: {str(e)}"

        # -------- Build Sources --------
        sources = []
        seen_sources = set()
        for d in docs:
            source_key = (d.metadata.get("source", "Unknown"), d.metadata.get("page", "N/A"))
            if source_key not in seen_sources:
                seen_sources.add(source_key)
                sources.append({
                    "document": d.metadata.get("source", "Unknown"),
                    "page": d.metadata.get("page", "N/A"),
                    "snippet": d.page_content[:200] + "..." if len(d.page_content) > 200 else d.page_content,
                })

        return {"answer": answer, "sources": sources}
